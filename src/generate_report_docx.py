r"""
Gera um arquivo .docx preenchendo seções do template com figuras e textos básicos.

Entrada esperada:
 - docs/report_template.md (opcional)
 - imagens em assets/figures/ e assets/interactive/ (PNG/HTML)

Saída:
 - output/dashboard_netflix.docx

Uso:
    python src\generate_report_docx.py
"""
import os
import sys
from docx import Document
from docx.shared import Inches

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
ASSETS = os.path.join(ROOT, 'assets')
FIGS = os.path.join(ASSETS, 'figures')
OUT = os.path.join(ROOT, 'output')
DOCX_OUT = os.path.join(OUT, 'dashboard_netflix.docx')

os.makedirs(OUT, exist_ok=True)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text):
    doc.add_paragraph(text)


def add_image_if_exists(doc, path, width_inches=6):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_inches))
    else:
        doc.add_paragraph(f'Imagem nao encontrada: {path}')


def generate():
    doc = Document()
    doc.add_heading('Relatório Final — Netflix Catalog Analysis', 0)
    add_paragraph(doc, 'Autor: \nData: ')

    add_heading(doc, '1. Introdução', level=1)
    add_paragraph(doc, 'Apresente a Netflix e o dataset (netflix_titles_CLEANED.csv).')

    add_heading(doc, '2. Metodologia', level=1)
    add_paragraph(doc, 'Descrever o processo de limpeza (preprocess.py), padronizacao de países e geracao de arquivos para Power BI.')

    add_heading(doc, '3. Resultados', level=1)
    add_heading(doc, 'Página 1 — Caracterização e Evolução do Catálogo', level=2)
    add_image_if_exists(doc, os.path.join(FIGS, 'page1_placeholder.png'))

    add_heading(doc, 'Página 2 — Estratégia de Produção Global (Top10)', level=2)
    add_image_if_exists(doc, os.path.join(FIGS, 'top10_countries.png'))

    add_heading(doc, 'Página 3 — Sazonalidade e Tendências', level=2)
    add_image_if_exists(doc, os.path.join(FIGS, 'heatmap_placeholder.png'))

    add_heading(doc, 'Página 4 — Análise de Duração e Classificação', level=2)
    add_image_if_exists(doc, os.path.join(FIGS, 'boxplot_duration_by_rating.png'))
    add_image_if_exists(doc, os.path.join(FIGS, 'bar_duration_series.png'))

    add_heading(doc, '4. Discussão', level=1)
    add_paragraph(doc, 'Discuta insights e relacoes entre as RQs.')

    add_heading(doc, '5. Limitações', level=1)
    add_paragraph(doc, 'Limitações do dataset e do método.')

    add_heading(doc, '6. Conclusão', level=1)
    add_paragraph(doc, 'Resumo dos principais achados.')

    doc.save(DOCX_OUT)
    print('Relatório gerado:', DOCX_OUT)
    return DOCX_OUT


if __name__ == '__main__':
    generate()
